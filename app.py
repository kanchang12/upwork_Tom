from flask import Flask, request, jsonify, render_template
import dropbox
from io import BytesIO
import docx  # Use `python-docx` package
import requests


app = Flask(__name__)

# Dropbox access token and file path
DROPBOX_ACCESS_TOKEN = 'YOUR_DROPBOX_ACCESS_TOKEN'
DROPBOX_FILE_PATH = '/path/to/your/document.docx'

# Initialize Dropbox client
dbx = dropbox.Dropbox(DROPBOX_ACCESS_TOKEN)

# Hard-coded file names and paths
DOCUMENTS = {
    "rules.docx": "https://www.dropbox.com/scl/fi/xuc0nwgoxjpvh02zlj6j6/rules.docx?rlkey=8udh1aziazvmy4zevyaojhu1w&st=kx7gtc5y&dl=0",
    "tasks.docx": "https://www.dropbox.com/scl/fi/k8cabdo6m35awe4ls0hi6/tasks.docx?rlkey=mklplg356j2cn8r457kiwh277&st=89yjka4t&dl=0"
}

# Claude AI API endpoint and headers
CLAUDE_API_URL = 'https://api.claude.ai/your_endpoint'
CLAUDE_API_HEADERS = {
    'Authorization': 'Bearer YOUR_CLAUDE_API_KEY',
    'Content-Type': 'application/json'
}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    user_input = request.json.get('user_input')
    
    # Send the command to Claude AI
    response = requests.post(CLAUDE_API_URL, headers=CLAUDE_API_HEADERS, json={'input': user_input})
    response_data = response.json()
    
    # Process the response from Claude AI
    action = response_data['action']
    args = response_data['args']
    
    if action == 'FETCH RECORDS':
        result = fetch_record(args['file_name'], args['variable_name'])
    elif action == 'UPDATE DOCUMENT':
        result = update_document(args['file_name'], args['variable_name'], args['new_value'])
    elif action == 'ADD NEW RECORD':
        result = add_new_record(args['file_name'], args['record'])
    elif action == 'ADD ALERT':
        result = add_alert(args['alert'])
    elif action == 'EXECUTE ALERT':
        result = execute_alert(args['alert_id'])
    else:
        result = 'Unknown action'
    
    return jsonify({'result': result})

def fetch_record(file_name, variable_name):
    # Read the document from Dropbox
    _, res = dbx.files_download(f'/path/to/your/{file_name}')
    doc = docx.Document(BytesIO(res.content))
    
    # Search for the variable name in the document
    for paragraph in doc.paragraphs:
        if variable_name in paragraph.text:
            return paragraph.text
    
    return f'{variable_name} not found in {file_name}'

def update_document(file_name, variable_name, new_value):
    # Read the document from Dropbox
    _, res = dbx.files_download(f'/path/to/your/{file_name}')
    doc = docx.Document(BytesIO(res.content))
    
    # Update the document with the new value
    for paragraph in doc.paragraphs:
        if variable_name in paragraph.text:
            paragraph.text = paragraph.text.replace(variable_name, new_value)
            break
    
    # Save the updated document back to Dropbox
    updated_content = doc_to_bytes(doc)
    dbx.files_upload(updated_content, f'/path/to/your/{file_name}', mode=dropbox.files.WriteMode.overwrite)
    
    return f'Document {file_name} updated with {new_value}'

def add_new_record(file_name, record):
    # Read the document from Dropbox
    _, res = dbx.files_download(f'/path/to/your/{file_name}')
    doc = docx.Document(BytesIO(res.content))
    
    # Add the new record
    doc.add_paragraph(record)
    
    # Save the updated document back to Dropbox
    updated_content = doc_to_bytes(doc)
    dbx.files_upload(updated_content, f'/path/to/your/{file_name}', mode=dropbox.files.WriteMode.overwrite)
    
    return f'New record added to {file_name}'

def add_alert(alert):
    # Placeholder: Implement your alert logic
    return f'Alert added: {alert}'

def execute_alert(alert_id):
    # Placeholder: Implement your alert execution logic
    return f'Alert executed: {alert_id}'

def doc_to_bytes(doc):
    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream.read()

if __name__ == '__main__':
    app.run(debug=True)
